"""
Script para generar la documentación técnica del proyecto
DIAN Auditor B2B en formato Word (.docx)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_bg(cell, hex_color):
    """Pinta el fondo de una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color_hex=None):
    p = doc.add_heading(text, level=level)
    if color_hex:
        for run in p.runs:
            run.font.color.rgb = RGBColor.from_string(color_hex)
    return p

def add_code_block(doc, code_text):
    """Añade un bloque de código con fuente monoespaciada y fondo gris."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    # Fondo gris
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    return p

def add_table(doc, headers, rows, header_color='1E3A5F'):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light List Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Encabezados
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].font.bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr[i], header_color)
    
    # Datos
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = val
            if r_idx % 2 == 0:
                set_cell_bg(row_cells[c_idx], 'EAF0FB')
    return table

def main():
    doc = Document()

    # Configurar márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # ════════════════════════════════════════
    #  PORTADA
    # ════════════════════════════════════════
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('DIAN Auditor B2B')
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)  # azul

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = subtitle.add_run('Documentación Técnica Completa del Sistema')
    sub.font.size = Pt(16)
    sub.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.add_run('Versión: 1.0  |  Autor: Santiago Parra  |  Fecha: Marzo 2026').font.size = Pt(10)

    doc.add_page_break()

    # ════════════════════════════════════════
    #  1. DESCRIPCIÓN GENERAL
    # ════════════════════════════════════════
    add_heading(doc, '1. Descripción General del Proyecto', level=1, color_hex='1E3A8A')
    doc.add_paragraph(
        'El DIAN Auditor B2B es un sistema Micro-SaaS de inteligencia de cumplimiento tributario '
        'desarrollado en Python con una interfaz web moderna. Está diseñado para empresas que '
        'necesitan auditar masivamente sus proveedores contra bases de datos gubernamentales colombianas, '
        'y para agencias de aduanas que requieren clasificar automáticamente sus catálogos de importación '
        'con códigos arancelarios (HS Codes).'
    )
    doc.add_paragraph(
        'El sistema integra tres grandes capacidades:\n'
        '1. Auditoría de riesgo tributario de proveedores (DIAN + BDME)\n'
        '2. Clasificación arancelaria masiva potenciada por IA local (Phi-4 con RAG)\n'
        '3. API REST de interoperabilidad para integración con ERPs externos'
    )

    # ════════════════════════════════════════
    #  2. ESTRUCTURA DEL PROYECTO
    # ════════════════════════════════════════
    add_heading(doc, '2. Estructura de Carpetas', level=1, color_hex='1E3A8A')
    doc.add_paragraph(
        'El proyecto está organizado de forma modular, siguiendo las convenciones del ecosistema '
        'Python para proyectos de ingeniería de datos y microservicios.'
    )

    add_table(doc,
        ['Carpeta / Archivo', 'Propósito'],
        [
            ['api/',             'Orquestador FastAPI. Expone los endpoints REST hacia el frontend y ERPs externos.'],
            ['src/',             'Núcleo de lógica de negocio. Módulos Python independientes por responsabilidad.'],
            ['data/',            'Datasets de entrada y caché de resultados (Parquet + CSV).'],
            ['data/raw/',        'Archivos originales como Proveedores-Ficticios-16022026.xlsx de la DIAN.'],
            ['data/samples/',    'Archivos de prueba (CSV de productos, Excel de facturación simulada).'],
            ['docs/',            'Documentación técnica y notas de arquitectura (PDR, Prompts).'],
            ['outputs/',         'Reportes Excel generados con codificación de color por nivel de riesgo.'],
            ['scripts/',         'Herramientas internas de depuración y utilidades de desarrollo.'],
            ['tests/',           'Suite de pruebas automatizadas con Pytest.'],
            ['logs/',            'Trazas de ejecución y archivos de salida textual.'],
            ['dashboard.html',   'Interfaz principal del Portal de Auditoría B2B.'],
            ['clasificador.html','Interfaz del Auto-Clasificador Arancelario IA.'],
        ]
    )

    # ════════════════════════════════════════
    #  3. MÓDULOS DEL BACKEND
    # ════════════════════════════════════════
    add_heading(doc, '3. Módulos del Backend (src/)', level=1, color_hex='1E3A8A')

    # 3.1 utils.py
    add_heading(doc, '3.1 utils.py — Normalización de NITs', level=2)
    doc.add_paragraph(
        'Contiene la función normalize_nit(raw_nit) que convierte cualquier formato de NIT colombiano '
        '(con puntos, guiones, espacios, como entero o float) a un string normalizado de exactamente 9 dígitos. '
        'Elimina el dígito verificador (separado por guión) y rellena con ceros a la izquierda si es necesario. '
        'Es usada por todos los demás módulos antes de cualquier operación de cruce de datos.'
    )
    add_code_block(doc,
        'normalize_nit("900.123.456-7")  →  "900123456"\n'
        'normalize_nit("12345678")         →  "012345678"\n'
        'normalize_nit(900123456)          →  "900123456"'
    )

    # 3.2 dian_processor.py
    add_heading(doc, '3.2 dian_processor.py — Procesador DIAN', level=2)
    doc.add_paragraph(
        'Implementado con Polars para permitir procesamiento out-of-core de archivos Excel de gran tamaño '
        '(e.g. "Proveedores Ficticios" con cientos de miles de registros). El módulo realiza:'
    )
    items = doc.add_paragraph(style='List Bullet')
    items.add_run('Lectura del Excel con Polars (calamine engine, sin cargar todo en RAM).')
    items = doc.add_paragraph(style='List Bullet')
    items.add_run('Validación del esquema (exige columna NIT).')
    items = doc.add_paragraph(style='List Bullet')
    items.add_run('Normalización de NITs con map_elements.')
    items = doc.add_paragraph(style='List Bullet')
    items.add_run('Deduplicación y adición de FECHA_INGESTA.')
    items = doc.add_paragraph(style='List Bullet')
    items.add_run('Escritura a dian_ficticios.parquet para búsquedas ultrarrápidas posteriores.')
    doc.add_paragraph(
        'La función check_nit_dian(nit) usa scan_parquet (lazy) para buscar un NIT específico '
        'sin cargar todo el archivo, retornando True si el NIT aparece en la lista de ficticios.'
    )

    # 3.3 bdme_scraper.py
    add_heading(doc, '3.3 bdme_scraper.py — Scraper BDME/CHIP', level=2)
    doc.add_paragraph(
        'Implementa un scraper Selenium en modo headless para consultar el Boletín de Deudores Morosos '
        'del Estado (BDME) del portal CHIP (Contaduría General de la Nación). Características clave:'
    )
    items = doc.add_paragraph(style='List Bullet')
    items.add_run('Función consult_nit_bdme(nit): navega al portal, busca el NIT y lee el estado (SIN_DEUDA / EN_MORA / INDETERMINADO).')
    items = doc.add_paragraph(style='List Bullet')
    items.add_run('Rate limiting: espera aleatoria de 1.5–3 segundos entre consultas para respetar los límites del gobierno.')
    items = doc.add_paragraph(style='List Bullet')
    items.add_run('Caché inteligente: almacena los resultados en un archivo Parquet (bdme_cache.parquet) con expiración de 30 días.')
    items = doc.add_paragraph(style='List Bullet')
    items.add_run('Procesamiento por lotes: consult_batch_bdme() acepta hasta 500 NITs por corrida, verificando la caché primero.')
    items = doc.add_paragraph(style='List Bullet')
    items.add_run('Resilencia: maneja errores de DOM y CAPTCHA retornando INDETERMINADO en vez de fallar completamente.')

    # 3.4 report_engine.py
    add_heading(doc, '3.4 report_engine.py — Motor de Reportes', level=2)
    doc.add_paragraph(
        'Motor principal de la auditoría de cumplimiento. Orquesta todos los módulos anteriores y genera '
        'reportes Excel con codificación de color. El flujo de generate_report(file_path) es:'
    )
    add_code_block(doc,
        'CSV/Excel del Cliente\n'
        '  ↓ Detección automática de columna NIT\n'
        '  ↓ Normalización robusta (safe_normalize)\n'
        '  ↓ Cruce DIAN Ficticios (check_nit_dian)\n'
        '  ↓ Cruce BDME (consult_batch_bdme)\n'
        '  ↓ Verificación facturador electrónico\n'
        '  ↓ Cálculo nivel de riesgo (ALTO / MEDIO / REVISAR / BAJO)\n'
        '  ↓ Excel con pestaña Detalle + pestaña Resumen ejecutivo\n'
        '     └─ Celdas coloreadas (Rojo / Naranja / Amarillo / Verde)'
    )
    add_table(doc,
        ['Nivel de Riesgo', 'Condición', 'Color'],
        [
            ['🔴 ALTO',   'NIT en lista DIAN ficticios',                        'Rojo'],
            ['🟠 MEDIO',  'En BDME (mora Estado) pero no en ficticios',          'Naranja'],
            ['🟡 REVISAR','Facturador no habilitado o estado indeterminado',     'Amarillo'],
            ['🟢 BAJO',   'Sin hallazgos en ninguna base de datos',              'Verde'],
        ]
    )

    # 3.5 clasificador.py
    add_heading(doc, '3.5 clasificador.py — Clasificador Arancelario IA con RAG', level=2)
    doc.add_paragraph(
        'Es el módulo más avanzado del sistema. Implementa una arquitectura RAG '
        '(Retrieval-Augmented Generation) para clasificar descripciones de mercancías en '
        'códigos HS (Harmonized System) de 10 dígitos con alta precisión. '
        'El modelo de IA utilizado es phi4:latest, ejecutado localmente vía Ollama (sin necesidad de internet).'
    )

    add_heading(doc, 'Cómo funciona la Arquitectura RAG:', level=3)
    add_code_block(doc,
        'PASO 1 — Vectorización del histórico:\n'
        '  data/historico_dian.csv (descripciones + HS Codes reales)\n'
        '   ↓ SentenceTransformer("all-MiniLM-L6-v2")\n'
        '   ↓ Vectores densos de 384 dimensiones\n'
        '   ↓ Indexados en FAISS / scikit-learn (en memoria)\n\n'
        'PASO 2 — Recuperación para nuevo producto:\n'
        '  "Máquina de café espresso industrial 220V"\n'
        '   ↓ Vector de embeddings\n'
        '   ↓ Búsqueda de similitud coseno en el índice\n'
        '   ↓ Top-3 ejemplos históricos más parecidos (score > 0.3)\n\n'
        'PASO 3 — Prompt aumentado a phi4:latest:\n'
        '  [SYSTEM PROMPT estricto] +\n'
        '  [Contexto histórico] +\n'
        '  [Descripción del producto]\n'
        '   ↓ phi4 genera JSON: {hs_code, confianza}\n'
        '   ↓ Parseo y validación robusta\n'
        '   ↓ Concatenado al DataFrame de resultados'
    )
    doc.add_paragraph(
        'Antes de RAG, el modelo asignaba capítulos incorrectos '
        '(p.ej. máquinas de café → cap. 09 "café en grano"; champú → cap. 15 "aceites"). '
        'Con RAG, el modelo ancla su respuesta en clasificaciones previamente validadas, '
        'eliminando los falsos positivos y el colapso de capítulos.'
    )
    add_table(doc,
        ['Producto', 'Sin RAG (Incorrecto)', 'Con RAG (Correcto)'],
        [
            ['Máquina de café espresso', '0901.11.00.00 (Café grano)', '8419.81.00.00 (Maq. industrial) ✅'],
            ['Champú de argán',          '1509.10.00.00 (Aceite oliva)','3305.10.00.00 (Champú) ✅'],
            ['Sillas de madera',         '6109.10 (Ropa de punto)',     '9401.69.00.00 (Muebles) ✅'],
            ['Smartphones 256GB',        '7000.xx.xx (Vidrio?)',        '8517.12.00.00 (Teléfonos) ✅'],
        ]
    )

    # ════════════════════════════════════════
    #  4. API FastAPI
    # ════════════════════════════════════════
    add_heading(doc, '4. Capa de Interoperabilidad (api/main.py)', level=1, color_hex='1E3A8A')
    doc.add_paragraph(
        'El archivo api/main.py implementa una API REST asíncrona con FastAPI que actúa como '
        'puente entre el frontend web y los módulos de backend. Características:'
    )
    add_table(doc,
        ['Endpoint', 'Método', 'Función'],
        [
            ['/',                     'GET',  'Health check. Devuelve información del sistema.'],
            ['/api/auditar-terceros', 'POST', 'Recibe archivo Excel/CSV, ejecuta la auditoría completa y devuelve el reporte coloreado.'],
            ['/api/clasificar-masivo','POST', 'Recibe CSV de productos, ejecuta el clasificador RAG y devuelve JSON con HS Codes.'],
        ]
    )
    doc.add_paragraph(
        'La API genera documentación Swagger interactiva automáticamente en http://localhost:8000/docs, '
        'lo que facilita la integración con ERPs externos. CORS habilitado para desarrollo.'
    )

    # ════════════════════════════════════════
    #  5. INTERFAZ WEB
    # ════════════════════════════════════════
    add_heading(doc, '5. Interfaces Web (Frontend)', level=1, color_hex='1E3A8A')

    add_heading(doc, '5.1 dashboard.html — Portal de Auditoría', level=2)
    doc.add_paragraph(
        'Interfaz principal del sistema. Es un Single Page Application (SPA) en HTML/CSS/JS Vanilla '
        '(sin frameworks externos). Incluye un diseño premium Dark Mode con:'
    )
    for item in [
        'Calculadora de riesgo interactiva por proveedor.',
        'Buscador bidireccional de HS Codes y nombres de productos.',
        'KPIs de cumplimiento con animaciones.',
        'Subida de archivos Excel/CSV para auditoría masiva.'
    ]:
        doc.add_paragraph(item, style='List Bullet')

    add_heading(doc, '5.2 clasificador.html — Portal IA', level=2)
    doc.add_paragraph(
        'Portal dedicado para el clasificador arancelario. Permite subir un CSV de productos '
        'y ver los resultados de la clasificación en tiempo real con una barra de progreso animada. '
        'Muestra el modelo IA activo (phi4:latest), niveles de confianza con badges de color, '
        'e incluye exportación a CSV de los resultados.'
    )

    # ════════════════════════════════════════
    #  6. DATOS Y CACHÉ
    # ════════════════════════════════════════
    add_heading(doc, '6. Gestión de Datos', level=1, color_hex='1E3A8A')
    add_table(doc,
        ['Archivo', 'Formato', 'Descripción'],
        [
            ['historico_dian.csv',       'CSV',     'Base de conocimiento RAG. 50+ clasificaciones arancelarias validadas por la DIAN.'],
            ['dian_ficticios.parquet',   'Parquet', 'Base de Proveedores Ficticios de la DIAN, procesada con Polars y normalizada.'],
            ['bdme_cache.parquet',       'Parquet', 'Caché de consultas BDME. Expira cada 30 días por NIT.'],
            ['outputs/*.xlsx',           'Excel',   'Reportes de cumplimiento generados automáticamente por report_engine.py.'],
        ]
    )

    # ════════════════════════════════════════
    #  7. TECNOLOGÍAS
    # ════════════════════════════════════════
    add_heading(doc, '7. Stack Tecnológico', level=1, color_hex='1E3A8A')
    add_table(doc,
        ['Tecnología', 'Versión', 'Uso en el Proyecto'],
        [
            ['Python',               '3.10+',   'Lenguaje principal del backend.'],
            ['FastAPI',              'última',   'API REST asíncrona y documentación Swagger.'],
            ['Polars',               'última',   'Procesamiento out-of-core de archivos DIAN masivos.'],
            ['Pandas',               '2.0+',     'Manipulación de DataFrames intermedios y exportación.'],
            ['Selenium',             '4.15+',    'Scraping web headless del portal CHIP/BDME.'],
            ['Ollama + phi4:latest', 'local',    'Motor de razonamiento IA sin internet.'],
            ['SentenceTransformers', 'última',   'Generación de embeddings para RAG (all-MiniLM-L6-v2).'],
            ['FAISS / scikit-learn', 'última',   'Búsqueda de similitud vectorial en memoria.'],
            ['openpyxl',             '3.1+',     'Escritura de archivos Excel con formato condicional.'],
            ['HTML / CSS / JS',      'Vanilla',  'Frontend sin dependencias externas.'],
        ]
    )

    # ════════════════════════════════════════
    #  8. INSTALACIÓN Y EJECUCIÓN
    # ════════════════════════════════════════
    add_heading(doc, '8. Instalación y Ejecución', level=1, color_hex='1E3A8A')

    add_heading(doc, 'Paso 1: Clonar y preparar entorno', level=2)
    add_code_block(doc,
        'git clone <url-del-repositorio>\n'
        'cd DIAN\n'
        'python -m venv .venv\n'
        '.venv\\Scripts\\activate   # Windows\n'
        'pip install -r requirements.txt'
    )

    add_heading(doc, 'Paso 2: Configurar modelo IA', level=2)
    doc.add_paragraph('Instalar Ollama desde https://ollama.com y descargar el modelo:')
    add_code_block(doc, 'ollama run phi4')

    add_heading(doc, 'Paso 3: Preparar la base DIAN', level=2)
    doc.add_paragraph('Coloca el archivo "Proveedores-Ficticios-DDMMYYYY.xlsx" en data/raw/ y ejecuta:')
    add_code_block(doc, 'python src/dian_processor.py')

    add_heading(doc, 'Paso 4: Iniciar el servidor API', level=2)
    add_code_block(doc,
        'python api/main.py\n'
        '# Swagger disponible en: http://localhost:8000/docs'
    )

    add_heading(doc, 'Paso 5: Abrir la interfaz', level=2)
    doc.add_paragraph('Abre dashboard.html o clasificador.html en el navegador. Los portales se comunican con la API en localhost:8000.')

    # ════════════════════════════════════════
    #  9. PRUEBAS
    # ════════════════════════════════════════
    add_heading(doc, '9. Pruebas Automatizadas', level=1, color_hex='1E3A8A')
    doc.add_paragraph('El proyecto incluye un suite Pytest ubicado en tests/:')
    add_code_block(doc, 'pytest tests/ -v')
    doc.add_paragraph(
        'Las pruebas cubren normalización de NITs, el procesador DIAN y la lógica de niveles de riesgo. '
        'Para las pruebas que usan Polars, asegurarse de que polars esté instalado en el entorno activo.'
    )

    # ════════════════════════════════════════
    #  GUARDAR
    # ════════════════════════════════════════
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '../docs/Documentacion_DIAN_Auditor_B2B.docx')
    doc.save(output_path)
    print(f"✅ Documento guardado en: {output_path}")

if __name__ == '__main__':
    main()
